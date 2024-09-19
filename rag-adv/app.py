import io

import streamlit as st
import faiss
import os
import numpy as np
from docx import Document
from PyPDF2 import PdfReader
from langchain.chains import RetrievalQA
from langchain_community.document_loaders import WebBaseLoader
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.tools.e2b_data_analysis.tool import UploadedFile
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.docstore.in_memory import InMemoryDocstore
from langchain_huggingface import HuggingFaceEndpoint

from dotenv import load_dotenv
load_dotenv()


def process_input(input_type, input_data):
    """
    Process Different Input Types and return a vector store
    :param input_type: type of input - pdf,link,doc,text
    :param input_data: data file for pdf/doc or link or text
    :return: vector store
    """
    global documents
    loader=None

    if input_type == 'links':
        loader = WebBaseLoader(input_data)
        documents = loader.load()
    elif input_type == 'Text':
        if isinstance(input_data, str):
            documents = input_data
        else:
            raise ValueError('Unsupported input type, String Expected')

    elif input_type == 'PDFs':
        if isinstance(input_data, io.BytesIO):
            pdf_reader = PdfReader(input_data)
        elif isinstance(input_data, UploadedFile):
            pdf_reader = PdfReader(io.BytesIO(input_data.read()))
        else:
            raise ValueError('Unsupported input type')
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        documents = text

    elif input_type == 'Docs':
        if isinstance(input_data, io.BytesIO):
            doc = Document(input_data)
        elif isinstance(input_data, UploadedFile):
            doc = Document(io.BytesIO(input_data.read()))
        else:
            raise ValueError('Unsupported input type')
        text = "\n".join(para.text for para in doc.paragraphs)
        documents = text

    elif input_type == 'TXTs':
        if isinstance(input_data, io.BytesIO):
            text = input_data.read().decode('utf-8')
        elif isinstance(input_data, UploadedFile):
            text = str(input_data.read().decode('utf-8'))
        else:
            raise ValueError('Unsupported input type')

    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    if input_type == 'Links':
        texts = text_splitter.split_documents(documents)
        texts = [str(doc.page_content) for doc in texts]
    else:
        texts = text_splitter.split_text(documents)

    model_name = "sentence-transformers/all-mpnet-base-v2"
    model_kwargs = {"device": "cpu"}
    encode_kwargs = {"normalize_embeddings": False}

    hf_embeddings = HuggingFaceEmbeddings(model_name=model_name, model_kwargs=model_kwargs, encode_kwargs=encode_kwargs)

    sample_embedding = np.array(hf_embeddings.embed_query("Sample text"))
    dimension = sample_embedding.shape[0]
    index = faiss.IndexFlatL2(dimension)

    vector_store = FAISS(
        embedding_function=hf_embeddings.embed_query,
        index=index,
        docstore=InMemoryDocstore(),
        index_to_docstore_id={}
    )
    vector_store.add_texts(texts)
    return vector_store


def answer_questions(vector_store,query):
    """
    Answer the questions based on provided input
    :param vector_store: vector store
    :param query: question
    :return: response
    """
    llm = HuggingFaceEndpoint(repo_id="meta-llama/Meta-Llama-3.1-8B-Instruct",
                              token=os.getenv("HUGGINGFACEHUB_API_TOKEN"), temperature=0.5)
    qa=RetrievalQA.from_chain_type(llm=llm, retriever=vector_store.as_retriever)
    answer = qa({"query": query})
    return answer



def main():
    st.text("Advance RAG and Q&A System")
    input_type = st.selectbox("Input type",["Links","PDFs","Text","Docs","TXTs"])
    if input_type == "Links":
        number_input = st.number_input(min_value=1,max_value=20,step=1,label="Enter number of links")
        input_data=[]
        for i in range(number_input):
            url = st.sidebar.text_input(f"URL {i+1}")
            input_data.append(url)
    elif input_type == "Text":
        input_data = st.text_input("Enter text to process")
    elif input_type == "PDFs":
        input_data = st.file_uploader("Upload PDF File", type=["pdf"])
    elif input_type == "Docs":
        input_data = st.file_uploader("Upload Docs File", type=["doc","docx"])
    elif input_type == "TXTs":
        input_data = st.file_uploader("Upload TXT File", type=["txt"])

    if st.button("Proceed"):
        try:
            vectorstore = process_input(input_type,input_data)
            st.session_state["vectorstore"] = vectorstore
            if "vectorstore" in st.session_state:
                query = st.text_input("Ask your question")
                if st.button("Submit"):
                    answer = answer_questions(st.session_state["vectorstore"],query)
                    st.write(answer)

        except Exception as e:
            st.error(f"ERROR IN SUBMISSION \n {e}")



if __name__ == "__main__":
    main()